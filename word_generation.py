# =============================================================================
# WORD DOCUMENT GENERATION
# =============================================================================
import os
import logging
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import re

# Configure logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def create_word_report(client_name, province, calculation_details, present_value_details, 
                      result, collateral_benefits, missed_time_unit, missed_time, output_path, 
                      birthdate=None, retirement_age=None, **kwargs):
    """
    Create a Word document report based on the template, filling in dynamic fields.
    
    Args:
        client_name: Name of the client
        province: Client's province for taxation purposes
        calculation_details: Dictionary containing past lost wages calculation details
        present_value_details: Dictionary containing future lost wages calculation details
        result: Dictionary containing income and deduction details
        collateral_benefits: Dictionary containing collateral benefit details
        missed_time_unit: Unit of missed time (days, weeks, months, etc.)
        missed_time: Amount of missed time
        output_path: Path where the Word document will be saved
        birthdate: Client's date of birth (optional)
        retirement_age: Client's retirement age (optional)
        **kwargs: Additional keyword arguments including:
            - loss_date: Date of loss/accident
            - current_date: Current date
            - ei_days_remaining: Remaining EI days
            - missed_pay: Missed pay amount
            - net_past_lost_wages: Net past lost wages amount
            - return_status: Return to work status
            - end_date: Speculative return to work date
    
    Returns:
        Path to the created Word document file
    """
    # Get optional parameters
    loss_date = kwargs.get('loss_date')
    current_date = kwargs.get('current_date', datetime.date.today())
    ei_days_remaining = kwargs.get('ei_days_remaining', 0)
    missed_pay = kwargs.get('missed_pay', 0)
    net_past_lost_wages = kwargs.get('net_past_lost_wages', calculation_details.get('Original Past Lost Wages', 
                                    calculation_details.get('Base Amount', 0)))
    return_status = kwargs.get('return_status', "").lower()
    end_date = kwargs.get('end_date', None)
    
    # Log all parameters for debugging
    logger.debug(f"Creating Word report for client: {client_name}, province: {province}")
    logger.debug(f"Loss date: {loss_date}, Current date: {current_date}")
    logger.debug(f"Return status: {return_status}, End date: {end_date}")
    logger.debug(f"Calculation details: {calculation_details}")
    logger.debug(f"Present value details: {present_value_details}")
    
    # Open the template file
    template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Lost Wages Report for.docx')
    logger.debug(f"Using template at: {template_path}")
    
    try:
        doc = Document(template_path)
        logger.debug("Successfully opened template document")
    except Exception as e:
        logger.error(f"Error opening template: {e}")
        # If template cannot be opened, create a new document
        doc = Document()
        doc.add_heading('PAST AND FUTURE WAGE LOSS', 0)
        logger.debug("Created new blank document as fallback")
    
    # Format dates
    today_date = current_date.strftime("%B %d, %Y")
    loss_date_str = loss_date.strftime("%B %d, %Y") if loss_date else "Not specified"
    
    # Get formatted province name
    if province.lower() == "nova scotia":
        province_display = "Nova Scotia"
    elif province.lower() == "new brunswick":
        province_display = "New Brunswick" 
    elif province.lower() == "prince edward island":
        province_display = "Prince Edward Island"
    elif province.lower() == "newfoundland" or province.lower() == "newfoundland and labrador":
        province_display = "Newfoundland and Labrador"
    else:
        province_display = province.capitalize()
    
    # Create a comprehensive dictionary of replacements
    replacements = {
        # Header section
        "[CLIENT NAME]": client_name,
        "[Today's Date]": today_date,
        
        # Jurisdiction section
        "[PROVINCE]": province_display,
        "[TAX YEAR]": str(current_date.year),
        
        # Income section
        "[GROSS INCOME]": f"${result.get('Gross Income', 0):,.2f}",
        "[FEDERAL TAXES]": f"${result.get('Federal Tax', 0):,.2f}",
        "[PROVINCIAL TAXES]": f"${result.get(f'{province.capitalize()} Tax', 0):,.2f}",
        "[CPP AMOUNT]": f"${result.get('CPP Contribution', 0) + result.get('CPP2 Contribution', 0):,.2f}",
        "[EI AMOUNT]": f"${result.get('EI Contribution', 0):,.2f}",
        "[DEPENDENT DEDUCTION]": f"${result.get('Dependent Benefit', 0):,.2f}",
        "[NET INCOME]": f"${result.get('Net Pay (Provincially specific deductions for damages)', 0):,.2f}",
        
        # Past Lost Wages section
        "[TIME PERIOD FOR PAST LOST WAGES]": f"{missed_time} {missed_time_unit}",
        "[TOTAL MISSED INCOME BEFORE PJI AND COLLATERAL BENEFITS]": f"${missed_pay:,.2f}" if isinstance(missed_pay, (int, float)) else f"${net_past_lost_wages:,.2f}",
        "[TOTAL COLLATERAL BENEFITS RECEIVED TO DATE AMOUNT]": f"${collateral_benefits.get('Total Past Benefits', 0):,.2f}",
        "[DATE RANGE USED TO CALCULATE PJI]": f"{loss_date_str} to {today_date}",
        "[PJI Rate]": f"{calculation_details.get('PJI Rate', 2.5):.2f}%",
        "[PJI Amount]": f"${calculation_details.get('Interest Amount', 0):,.2f}",
        "[TOTAL PAST LOST WAGES AFTER PJI AND COLLATERAL BENEFITS]": f"${calculation_details.get('Past Lost Wages with Interest', 0):,.2f}",
    }
    
    # Add individual collateral benefits to relevant sections
    collateral_benefits_text = f"${collateral_benefits.get('Total Past Benefits', 0):,.2f}"
    
    if collateral_benefits.get('EI Benefits (to date)', 0) > 0:
        collateral_benefits_text += f"\nEI Benefits: ${collateral_benefits.get('EI Benefits (to date)', 0):,.2f}"
    if collateral_benefits.get('Section B Benefits (to date)', 0) > 0:
        collateral_benefits_text += f"\nSection B Benefits: ${collateral_benefits.get('Section B Benefits (to date)', 0):,.2f}"
    if collateral_benefits.get('LTD Benefits (to date)', 0) > 0:
        collateral_benefits_text += f"\nLTD Benefits: ${collateral_benefits.get('LTD Benefits (to date)', 0):,.2f}"
    if collateral_benefits.get('CPPD Benefits (to date)', 0) > 0:
        collateral_benefits_text += f"\nCPPD Benefits: ${collateral_benefits.get('CPPD Benefits (to date)', 0):,.2f}"
    if collateral_benefits.get('Other Benefits (to date)', 0) > 0:
        collateral_benefits_text += f"\nOther Benefits: ${collateral_benefits.get('Other Benefits (to date)', 0):,.2f}"
    
    replacements["[TOTAL COLLATERAL BENEFITS RECEIVED TO DATE AMOUNT]"] = collateral_benefits_text
    
    # Future Lost Wages section - only included if present value is positive
    calculate_future_wages = present_value_details.get('present_value', 0) > 0
    
    # Process conditional future wages section
    if calculate_future_wages:
        logger.debug("Including future lost wages section")
        
        future_replacements = {
            "[NET INCOME]": f"${result.get('Net Pay (Provincially specific deductions for damages)', 0):,.2f}",
            "[RETURN TO WORK STATUS]": (return_status.capitalize() if return_status else "Not specified"),
            "[DISCOUNT RATE PERCENTAGE]": f"{present_value_details.get('discount_rate', 0) * 100:.2f}%",
            "[Future Lost Wages Time Horizon]": f"{present_value_details.get('time_horizon', 0):.2f} years",
            "[TOTAL FUTURE LOST WAGES AMOUNT]": f"${present_value_details.get('present_value', 0):,.2f}"
        }
        
        # Add future collateral benefits
        future_benefits_text = f"${collateral_benefits.get('Total Annual Future Benefits', 0):,.2f}"
        
        if collateral_benefits.get('EI Benefits (annual)', 0) > 0:
            future_benefits_text += f"\nEI Benefits: ${collateral_benefits.get('EI Benefits (annual)', 0):,.2f}"
        if collateral_benefits.get('Section B Benefits (annual)', 0) > 0:
            future_benefits_text += f"\nSection B Benefits: ${collateral_benefits.get('Section B Benefits (annual)', 0):,.2f}"
        if collateral_benefits.get('LTD Benefits (annual)', 0) > 0:
            future_benefits_text += f"\nLTD Benefits: ${collateral_benefits.get('LTD Benefits (annual)', 0):,.2f}"
        if collateral_benefits.get('CPPD Benefits (annual)', 0) > 0:
            future_benefits_text += f"\nCPPD Benefits: ${collateral_benefits.get('CPPD Benefits (annual)', 0):,.2f}"
        if collateral_benefits.get('Other Benefits (annual)', 0) > 0:
            future_benefits_text += f"\nOther Benefits: ${collateral_benefits.get('Other Benefits (annual)', 0):,.2f}"
        
        future_replacements["[TOTAL Annual Collateral Benefits Moving Forward]"] = future_benefits_text
        
        # Handle return to work conditional sections
        if "returning to work" in return_status:
            logger.debug("Including 'returning to work' conditional section")
            formatted_end_date = "Not specified"
            
            if end_date:
                if hasattr(end_date, 'strftime'):
                    formatted_end_date = end_date.strftime("%B %d, %Y")
                else:
                    formatted_end_date = str(end_date)
                    
            future_replacements["[Speculative Return to Work Date]"] = formatted_end_date
            
        # Handle total disability conditional sections
        if "total disability" in return_status:
            logger.debug("Including 'total disability' conditional section")
            if birthdate:
                future_replacements["[Date of Birth]"] = birthdate.strftime("%B %d, %Y")
            else:
                future_replacements["[Date of Birth]"] = "Not specified"
                
            future_replacements["[Retirement Age]"] = str(retirement_age) if retirement_age else "65"
        
        # Add future lost wages replacements to main replacements dictionary
        replacements.update(future_replacements)
    
    # Total Wage Loss section
    replacements.update({
        "[TOTAL PAST LOST WAGES]": f"${calculation_details.get('Past Lost Wages with Interest', 0):,.2f}",
        "[TOTAL FUTURE LOST WAGES]": f"${present_value_details.get('present_value', 0):,.2f}",
        "[Total Economic Damages]": f"${calculation_details.get('Past Lost Wages with Interest', 0) + present_value_details.get('present_value', 0):,.2f}"
    })
    
    # Notes section
    provincial_notes = ""
    if province.lower() == "nova scotia":
        provincial_notes = "In Nova Scotia, CPP, EI, and income taxes are deducted when calculating net income for damages."
    elif province.lower() == "new brunswick":
        provincial_notes = "In New Brunswick, federal and provincial income taxes are deducted, but CPP and EI are NOT deducted for damages calculations."
    elif province.lower() == "prince edward island":
        provincial_notes = "PEI calculates damages based on gross income rather than net income."
    elif province.lower() == "newfoundland" or province.lower() == "newfoundland and labrador":
        provincial_notes = "Newfoundland follows the standard 'net income' approach with province-specific tax calculations."
    
    pji_note = f"PJI Rate was calculated using T-Bill rates for the period from {loss_date_str} to {today_date}"
    discount_rate_note = f"Discount rate of {present_value_details.get('discount_rate', 0) * 100:.2f}% is used as per {province_display} standards"
    
    # Add notes to replacements
    replacements["[NOTE on HOW PJI Rate Was calculate]"] = pji_note
    replacements["[NOTE ON ANY PROVICIALLY SPECIFIC REASONING ON DEDUCTIONS]"] = provincial_notes
    replacements["[NOTE ON HOW DISCOUNT RATE WAS CALCULATED]"] = discount_rate_note
    replacements["[NOTE ANY PROVINCIALLY SPECIFIC REASONING]"] = provincial_notes
    
    # EI benefits note
    if ei_days_remaining > 0:
        replacements["[NOTE EI SICK BENIFITS CUT OFF IF APPLICABLE]"] = f"EI sickness benefits are limited to 26 weeks (182 days). {ei_days_remaining} days remaining."
    else:
        replacements["[NOTE EI SICK BENIFITS CUT OFF IF APPLICABLE]"] = "EI sickness benefits period has been fully utilized."
    
    # Log the replacements for debugging
    logger.debug(f"Created {len(replacements)} replacement mappings")
    
    # Process document content
    logger.debug("Processing document content")
    
    # First, identify all conditional sections
    conditional_sections = []
    current_section = None
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        
        # Start of future lost wages section
        if "*ONLY SHOW THIS IF THEY ASKED TO CALCULATE FUTURE LOST WAGES*" in text:
            current_section = {
                'type': 'future_lost_wages',
                'start': i,
                'end': None,
                'include': calculate_future_wages
            }
            conditional_sections.append(current_section)
        
        # Start of return to work conditional section
        elif "*ONLY IF RETURN TO WORK STATUS IS Returning to Work*" in text:
            current_section = {
                'type': 'returning_to_work',
                'start': i,
                'end': None,
                'include': "returning to work" in return_status
            }
            conditional_sections.append(current_section)
        
        # Start of total disability conditional section
        elif "*ONLY IF RETURN TO WORK STATUS IS TOTAL DISABILITY*" in text:
            current_section = {
                'type': 'total_disability',
                'start': i,
                'end': None,
                'include': "total disability" in return_status
            }
            conditional_sections.append(current_section)
        
        # End of future lost wages section (find the total wage loss section)
        elif current_section and current_section['type'] == 'future_lost_wages' and "**TOTAL WAGE LOSS**" in text:
            current_section['end'] = i
            current_section = None
        
        # End of other conditional sections (end at next line or after processing)
        elif current_section and current_section['end'] is None:
            current_section['end'] = i + 1
    
    # Process regular replacements in the document
    for paragraph in doc.paragraphs:
        # Skip processing if this paragraph should be entirely removed
        skip = False
        for section in conditional_sections:
            if not section['include'] and section['start'] <= i < section['end']:
                skip = True
                break
        
        if skip:
            continue
            
        # Process replacements for this paragraph
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))
    
    # Also process tables for replacements
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value))
    
    # Handle conditional sections - remove conditional markers
    for i, paragraph in enumerate(doc.paragraphs):
        # Clean up conditional markers
        if "*ONLY SHOW THIS IF THEY ASKED TO CALCULATE FUTURE LOST WAGES*" in paragraph.text:
            if calculate_future_wages:
                paragraph.text = paragraph.text.replace("*ONLY SHOW THIS IF THEY ASKED TO CALCULATE FUTURE LOST WAGES*", "")
            else:
                paragraph.text = ""
                
        elif "*ONLY IF RETURN TO WORK STATUS IS Returning to Work*" in paragraph.text:
            if "returning to work" in return_status:
                paragraph.text = paragraph.text.replace("*ONLY IF RETURN TO WORK STATUS IS Returning to Work*", "")
            else:
                paragraph.text = ""
                
        elif "*ONLY IF RETURN TO WORK STATUS IS TOTAL DISABILITY*" in paragraph.text:
            if "total disability" in return_status:
                paragraph.text = paragraph.text.replace("*ONLY IF RETURN TO WORK STATUS IS TOTAL DISABILITY*", "")
            else:
                paragraph.text = ""
    
    # Save the document
    try:
        doc.save(output_path)
        logger.info(f"Word document saved successfully at: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"Error saving Word document: {e}")
        return None
